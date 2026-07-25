"""Resume DOCX generator using python-docx.

Generates a structured resume with work experience, education, and skills
from an ApplicantProfile. Falls back to python-docx when resumecraft is unavailable.
"""

import random
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches

from src.data_generation.profile import ApplicantProfile
from src.data_generation.utils import random_date_between


UAE_UNIVERSITIES: list[str] = [
    "UAE University",
    "Zayed University",
    "American University of Sharjah",
    "United Arab Emirates University",
    "Khalifa University",
    "American University in Dubai",
    "University of Sharjah",
    "Ajman University",
    "Abu Dhabi University",
    "Dubai Institute of Design and Innovation",
]

DEGREE_TYPES: list[str] = [
    "Bachelor",
    "Master",
    "Associate",
    "Diploma",
]

UAE_CITIES: list[str] = [
    "Abu Dhabi",
    "Dubai",
    "Sharjah",
    "Ajman",
    "Al Ain",
    "Ras Al Khaimah",
    "Fujairah",
]

OCCUPATION_SKILLS: dict[str, dict[str, list[str]]] = {
    "Software Engineer": {
        "Languages": ["Python", "JavaScript", "TypeScript", "Java", "Go", "Rust"],
        "Technical": ["Git", "Docker", "Kubernetes", "AWS", "PostgreSQL", "Redis", "CI/CD", "Linux"],
        "Soft Skills": ["Problem Solving", "Team Collaboration", "Communication", "Agile"],
    },
    "Accountant": {
        "Languages": [],
        "Technical": ["Excel", "QuickBooks", "SAP", "Xero", "Financial Reporting", "Tax Filing", "Audit"],
        "Soft Skills": ["Attention to Detail", "Analytical Thinking", "Time Management", "Integrity"],
    },
    "Teacher": {
        "Languages": ["Arabic", "English"],
        "Technical": ["Google Classroom", "Microsoft Teams", "Canvas LMS", "Smart Board"],
        "Soft Skills": ["Communication", "Patience", "Classroom Management", "Curriculum Design"],
    },
    "Nurse": {
        "Languages": ["Arabic", "English"],
        "Technical": ["Patient Care", "IV Therapy", "EMR Systems", "First Aid", "Vital Signs Monitoring"],
        "Soft Skills": ["Empathy", "Communication", "Critical Thinking", "Stress Management"],
    },
    "Sales Manager": {
        "Languages": ["Arabic", "English"],
        "Technical": ["CRM (Salesforce)", "Microsoft Office", "Data Analysis", "Market Research"],
        "Soft Skills": ["Negotiation", "Leadership", "Communication", "Relationship Building"],
    },
    "Driver": {
        "Languages": ["Arabic", "English", "Hindi", "Urdu"],
        "Technical": ["UAE Driving License", "GPS Navigation", "Vehicle Maintenance", "Defensive Driving"],
        "Soft Skills": ["Punctuality", "Reliability", "Customer Service", "Time Management"],
    },
    "Civil Engineer": {
        "Languages": ["Arabic", "English"],
        "Technical": ["AutoCAD", "Revit", "SAP2000", "Project Management", "Structural Analysis", "MS Project"],
        "Soft Skills": ["Problem Solving", "Team Leadership", "Communication", "Attention to Detail"],
    },
    "Administrative Assistant": {
        "Languages": ["Arabic", "English"],
        "Technical": ["Microsoft Office", "Google Workspace", "Data Entry", "Filing", "Scheduling"],
        "Soft Skills": ["Organization", "Communication", "Multitasking", "Attention to Detail"],
    },
    "Chef": {
        "Languages": ["Arabic", "English", "Hindi", "Urdu"],
        "Technical": ["Food Safety", "Menu Planning", "Kitchen Management", "Inventory Control", "Cuisine Specialization"],
        "Soft Skills": ["Creativity", "Teamwork", "Time Management", "Stress Management"],
    },
    "Electrician": {
        "Languages": ["Arabic", "English", "Hindi", "Urdu"],
        "Technical": ["Wiring", "Circuit Design", "Electrical Safety", "Blueprint Reading", "Troubleshooting"],
        "Soft Skills": ["Problem Solving", "Attention to Detail", "Physical Stamina", "Customer Service"],
    },
}

INDUSTRIES: list[str] = [
    "Technology",
    "Finance",
    "Healthcare",
    "Education",
    "Construction",
    "Retail",
    "Hospitality",
    "Manufacturing",
    "Government",
    "Transportation",
]


def _generate_company_name(rng: random.Random, exclude: str | None = None) -> str:
    """Generate a plausible UAE company name."""
    prefixes = [
        "Al Noor", "Emirates", "Gulf", "Arabian", "Dubai", "Abu Dhabi",
        "Sharjah", "National", "United", "Federal", "Crescent", "Pearl",
        "Golden", "Skyline", "Oasis", "Desert", "Falcon", "Palm",
    ]
    sectors = [
        "Trading", "General Trading", "Construction", "Services",
        "Technologies", "Solutions", "Enterprises", "Group",
        "Industries", "Logistics", "Real Estate", "Consulting",
        "Healthcare", "Education", "Finance", "Retail",
    ]
    suffixes = ["LLC", "LLC", "LLC", "LLC", "FZ-LLC", "FZE"]
    prefix = rng.choice(prefixes)
    sector = rng.choice(sectors)
    suffix = rng.choice(suffixes)
    name = f"{prefix} {sector} {suffix}"
    if exclude and name == exclude:
        return _generate_company_name(rng, exclude)
    return name


def _calculate_duration_months(start: date, end: date | None) -> int:
    """Calculate the duration in months between two dates."""
    if end is None:
        end = date.today()
    months = (end.year - start.year) * 12 + (end.month - start.month)
    if end.day < start.day:
        months -= 1
    return max(1, months)


def _generate_work_experience(
    profile: ApplicantProfile,
    num_positions: int,
    rng: random.Random,
) -> list[dict[str, Any]]:
    """Generate work experience records with cross-document consistency."""
    experiences: list[dict[str, Any]] = []
    today = date.today()

    # Generate past positions first (reverse chronological order in output)
    past_positions = num_positions - 1
    current_start = random_date_between(
        rng,
        profile.date_of_birth + timedelta(days=365 * 18),
        today - timedelta(days=30),
    )

    # Generate past experiences
    past_end = current_start
    for i in range(past_positions):
        job_title = profile.occupation
        company = _generate_company_name(rng, exclude=profile.employer_name)
        start = random_date_between(
            rng,
            profile.date_of_birth + timedelta(days=365 * 18),
            past_end - timedelta(days=60),
        )
        end = past_end
        location = rng.choice(UAE_CITIES)
        duration = _calculate_duration_months(start, end)

        achievements = _generate_achievements(profile.occupation, rng)

        experiences.append({
            "job_title": job_title,
            "company": company,
            "start_date": start.isoformat(),
            "location": location,
            "end_date": end.isoformat(),
            "is_current": False,
            "description": _generate_job_description(profile.occupation, rng),
            "achievements": achievements,
            "duration_months": duration,
            "industry": rng.choice(INDUSTRIES),
        })
        past_end = start

    # Current position (always last in list, marked as current)
    duration = _calculate_duration_months(current_start, None)
    achievements = _generate_achievements(profile.occupation, rng)

    experiences.append({
        "job_title": profile.occupation,
        "company": profile.employer_name,
        "start_date": current_start.isoformat(),
        "location": profile.address.get("city", rng.choice(UAE_CITIES)),
        "end_date": None,
        "is_current": True,
        "description": _generate_job_description(profile.occupation, rng),
        "achievements": achievements,
        "duration_months": duration,
        "industry": rng.choice(INDUSTRIES),
    })

    # Reverse to get chronological order (oldest first)
    experiences.reverse()
    return experiences


def _generate_achievements(occupation: str, rng: random.Random) -> list[str]:
    """Generate realistic job achievements for an occupation."""
    achievement_pool = {
        "Software Engineer": [
            "Led development of microservices architecture reducing latency by 40%",
            "Mentored junior developers and conducted code reviews",
            "Implemented CI/CD pipelines reducing deployment time by 60%",
            "Optimized database queries improving response time by 35%",
            "Designed and deployed RESTful APIs for mobile applications",
            "Contributed to open-source projects with 500+ GitHub stars",
        ],
        "Accountant": [
            "Reduced month-end closing time by 3 days through process automation",
            "Managed annual budget of AED 5M+ with 99% accuracy",
            "Implemented new accounting software reducing errors by 25%",
            "Led successful external audit with zero non-conformance findings",
            "Streamlined accounts payable process saving AED 50K annually",
        ],
        "Teacher": [
            "Developed curriculum adopted by 10+ schools district-wide",
            "Achieved 95% student pass rate in standardized exams",
            "Led after-school program with 80+ student enrollment",
            "Received Excellence in Teaching Award 2024",
            "Integrated technology into classroom improving engagement by 40%",
        ],
        "Nurse": [
            "Managed care for 30+ patients per shift in ICU unit",
            "Reduced patient readmission rates by 15% through care coordination",
            "Trained 20+ new nursing staff on EMR systems",
            "Received Patient Care Excellence Award",
            "Implemented infection control protocols reducing HAIs by 30%",
        ],
        "Sales Manager": [
            "Exceeded annual sales targets by 25% for 3 consecutive years",
            "Expanded client base by 40% through strategic partnerships",
            "Led team of 12 sales representatives achieving top regional performance",
            "Negotiated contracts worth AED 10M+ annually",
            "Implemented CRM system improving lead conversion by 35%",
        ],
        "Driver": [
            "Maintained 100% on-time delivery record for 2+ years",
            "Zero accidents over 100,000+ km driven",
            "Trained new drivers on safety protocols and route optimization",
            "Received Customer Service Excellence recognition",
            "Optimized delivery routes reducing fuel costs by 15%",
        ],
    }
    default_achievements = [
        "Consistently met performance targets and KPIs",
        "Recognized for outstanding contribution to team goals",
        "Led process improvement initiatives",
        "Mentored new team members",
        "Received Employee of the Month award",
    ]
    pool = achievement_pool.get(occupation, default_achievements)
    num = rng.randint(2, min(4, len(pool)))
    return rng.sample(pool, num)


def _generate_job_description(occupation: str, rng: random.Random) -> str:
    """Generate a brief job description for an occupation."""
    descriptions = {
        "Software Engineer": "Responsible for designing, developing, and maintaining software applications. Collaborated with cross-functional teams to deliver high-quality solutions.",
        "Accountant": "Managed financial records, prepared reports, and ensured compliance with regulatory requirements. Handled accounts payable/receivable and monthly reconciliations.",
        "Teacher": "Delivered curriculum-aligned instruction to students. Created engaging lesson plans and assessed student progress through various evaluation methods.",
        "Nurse": "Provided direct patient care in clinical settings. Administered medications, monitored vital signs, and coordinated with medical teams for treatment plans.",
        "Sales Manager": "Led sales team to achieve revenue targets. Developed sales strategies, managed key accounts, and conducted market analysis for business growth.",
        "Driver": "Operated vehicles for passenger/ goods transportation. Ensured timely deliveries, maintained vehicle logs, and adhered to traffic safety regulations.",
        "Civil Engineer": "Designed and supervised construction projects. Prepared technical drawings, conducted site inspections, and ensured compliance with building codes.",
        "Administrative Assistant": "Provided administrative support to management and teams. Handled scheduling, correspondence, filing, and coordinated office operations.",
        "Chef": "Prepared and presented high-quality dishes. Managed kitchen operations, menu planning, inventory, and maintained food safety standards.",
        "Electrician": "Installed and maintained electrical systems. Performed troubleshooting, repairs, and ensured compliance with safety codes and regulations.",
    }
    return descriptions.get(occupation, f"Performed duties related to {occupation.lower()} role. Collaborated with team members to achieve departmental objectives.")


def _generate_education(
    profile: ApplicantProfile,
    rng: random.Random,
) -> list[dict[str, Any]]:
    """Generate 1-2 education entries with UAE universities."""
    num_entries = rng.randint(1, 2)
    entries: list[dict[str, Any]] = []

    dob = profile.date_of_birth
    today = date.today()

    for i in range(num_entries):
        university = rng.choice(UAE_UNIVERSITIES)
        degree = rng.choice(DEGREE_TYPES)
        field = rng.choice([
            "Computer Science", "Business Administration", "Accounting",
            "Engineering", "Education", "Nursing", "Marketing",
            "Finance", "Law", "Medicine", "Information Technology",
        ])
        grad_year_start = max(dob.year + 18, today.year - 15)
        grad_year_end = today.year - 1
        if grad_year_start > grad_year_end:
            grad_year_start = grad_year_end
        grad_year = rng.randint(grad_year_start, grad_year_end)
        start_year = grad_year - (4 if degree == "Bachelor" else 2 if degree == "Master" else 3)

        entries.append({
            "institution": university,
            "degree": degree,
            "field_of_study": field,
            "start_date": f"{start_year}-09-01",
            "end_date": f"{grad_year}-06-01",
            "location": rng.choice(UAE_CITIES),
            "gpa": f"{rng.uniform(2.5, 4.0):.2f}",
        })

    return entries


def _generate_skills(
    occupation: str,
    rng: random.Random,
) -> dict[str, list[str]]:
    """Generate skills based on occupation with category breakdown."""
    skill_map = OCCUPATION_SKILLS.get(occupation)
    if not skill_map:
        # Fallback generic skills
        return {
            "Languages": ["Arabic", "English"],
            "Technical": ["Microsoft Office", "Email", "Documentation"],
            "Soft Skills": ["Communication", "Teamwork", "Problem Solving"],
        }

    skills: dict[str, list[str]] = {}
    for category, pool in skill_map.items():
        if pool:
            num = rng.randint(max(1, len(pool) // 3), min(len(pool), 5))
            skills[category] = rng.sample(pool, num)
    return skills


def _generate_certifications(
    occupation: str,
    rng: random.Random,
) -> list[dict[str, Any]]:
    """Generate 0-3 professional certifications."""
    cert_pool = {
        "Software Engineer": [
            "AWS Certified Developer", "Oracle Java SE", "Microsoft Azure Fundamentals",
            "Certified Kubernetes Administrator", "PMP",
        ],
        "Accountant": [
            "CPA", "ACCA", "CMA", "QuickBooks Certified", "SAP Certified",
        ],
        "Teacher": [
            "Teaching License", "TESOL Certificate", "Google Educator", "Cambridge CELTA",
        ],
        "Nurse": [
            "BLS Certification", "ACLS Certification", "RN License", "Critical Care Nursing",
        ],
        "Sales Manager": [
            "Salesforce Certified", "HubSpot Inbound Sales", "PMP", "Digital Marketing",
        ],
        "Driver": [
            "UAE Driving License", "Defensive Driving Certificate", "First Aid Certified",
        ],
        "Civil Engineer": [
            "PE License", "LEED Accredited", "PMP", "AutoCAD Certified", "OSHA Safety",
        ],
    }
    default_certs = [
        "Professional Development Certificate", "Industry Certification",
    ]

    pool = cert_pool.get(occupation, default_certs)
    num = rng.randint(0, min(3, len(pool)))
    if num == 0:
        return []

    certs = rng.sample(pool, num)
    today = date.today()
    result = []
    for cert in certs:
        issue_date = random_date_between(
            rng,
            today - timedelta(days=365 * 8),
            today - timedelta(days=30),
        )
        expiry_years = rng.choice([2, 3, 5, None])
        expiry_date = None
        if expiry_years:
            expiry_date = (issue_date + timedelta(days=365 * expiry_years)).isoformat()
            if date.fromisoformat(expiry_date) < today:
                expiry_date = None  # expired certs shown as None

        result.append({
            "name": cert,
            "issuing_organization": rng.choice([
                "Professional Body", "Industry Authority", "Certification Board",
                "Government Authority", "International Body",
            ]),
            "issue_date": issue_date.isoformat(),
            "expiry_date": expiry_date,
        })
    return result


def _generate_summary(occupation: str, years_exp: int, rng: random.Random) -> str:
    """Generate a professional summary paragraph."""
    templates = [
        f"Experienced {occupation} with {years_exp}+ years of proven track record in delivering results. "
        "Known for strong work ethic, attention to detail, and ability to work effectively in diverse teams.",
        f"Dedicated {occupation} with {years_exp} years of experience in UAE market. "
        "Committed to excellence and continuous professional development.",
        f"Results-driven {occupation} bringing {years_exp} years of hands-on experience. "
        "Proven ability to adapt to new challenges and contribute to organizational success.",
    ]
    return rng.choice(templates)


def _build_resume_docx(
    resume_data: dict[str, Any],
    work_experiences: list[dict[str, Any]],
    education_entries: list[dict[str, Any]],
    skills: dict[str, list[str]],
    certifications: list[dict[str, Any]],
    output_path: Path,
) -> None:
    """Build a formatted DOCX resume using python-docx."""
    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # Name header
    name = resume_data.get("full_name", "Applicant Name")
    title = doc.add_heading(name, level=0)
    title.alignment = 1  # center

    # Contact info
    contact_parts = []
    if resume_data.get("email"):
        contact_parts.append(resume_data["email"])
    if resume_data.get("phone"):
        contact_parts.append(resume_data["phone"])
    if resume_data.get("location"):
        contact_parts.append(resume_data["location"])
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = 1
        contact_para.paragraph_format.space_after = Pt(6)

    # Summary
    if resume_data.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(resume_data["summary"])

    # Work Experience
    if work_experiences:
        doc.add_heading("Work Experience", level=1)
        for exp in work_experiences:
            # Job title and company on same line
            job_line = f"{exp['job_title']} — {exp['company']}"
            p = doc.add_paragraph()
            run = p.add_run(job_line)
            run.bold = True

            # Date and location
            if exp.get("end_date"):
                date_range = f"{exp['start_date']} – {exp['end_date']}"
            else:
                date_range = f"{exp['start_date']} – Present"
            location_info = exp.get("location", "")
            if location_info:
                date_range += f" | {location_info}"
            date_para = doc.add_paragraph(date_range)
            date_para.paragraph_format.space_after = Pt(2)

            # Description
            if exp.get("description"):
                doc.add_paragraph(exp["description"])

            # Achievements
            if exp.get("achievements"):
                for achievement in exp["achievements"]:
                    doc.add_paragraph(achievement, style="List Bullet")

            doc.add_paragraph()  # spacing

    # Education
    if education_entries:
        doc.add_heading("Education", level=1)
        for edu in education_entries:
            degree_line = f"{edu['degree']} in {edu['field_of_study']}"
            p = doc.add_paragraph()
            run = p.add_run(degree_line)
            run.bold = True

            school_line = edu["institution"]
            if edu.get("location"):
                school_line += f", {edu['location']}"
            if edu.get("start_date") and edu.get("end_date"):
                school_line += f" | {edu['start_date'][:4]} – {edu['end_date'][:4]}"
            doc.add_paragraph(school_line)

            if edu.get("gpa"):
                doc.add_paragraph(f"GPA: {edu['gpa']}")
            doc.add_paragraph()

    # Skills
    if skills:
        doc.add_heading("Skills", level=1)
        for category, skill_list in skills.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{category}: ")
            run.bold = True
            p.add_run(", ".join(skill_list))

    # Certifications
    if certifications:
        doc.add_heading("Certifications", level=1)
        for cert in certifications:
            p = doc.add_paragraph()
            run = p.add_run(cert["name"])
            run.bold = True
            extra = f" — {cert['issuing_organization']}"
            if cert.get("issue_date"):
                extra += f" ({cert['issue_date'][:4]}"
                if cert.get("expiry_date"):
                    extra += f" – {cert['expiry_date'][:4]}"
                extra += ")"
            p.add_run(extra)

    doc.save(output_path)


def generate_resume(
    profile: ApplicantProfile,
    seed: int,
) -> tuple[dict[str, Any], list[dict[str, Any]], Path]:
    """Generate a resume DOCX with structured data and work experience records.

    Args:
        profile: ApplicantProfile seed object for cross-document consistency.
        seed: Random seed for reproducibility.

    Returns:
        Tuple of (resume_data dict, work_experience list, DOCX file path).
        The resume_data dict maps to the resume_data schema (14 fields).
        The work_experience list maps to resume_work_experience schema (10 fields each).
    """
    rng = random.Random(seed)

    # Calculate years of experience based on age
    age = (date.today() - profile.date_of_birth).days // 365
    years_start_working = rng.randint(18, 22)
    years_of_experience = max(1, age - years_start_working)

    # Number of positions: 2-5
    num_positions = rng.randint(2, 5)

    # Generate work experience
    work_experiences = _generate_work_experience(profile, num_positions, rng)

    # Generate education
    education_entries = _generate_education(profile, rng)

    # Generate skills
    skills = _generate_skills(profile.occupation, rng)
    skill_count = sum(len(v) for v in skills.values())

    # Generate certifications
    certifications = _generate_certifications(profile.occupation, rng)

    # Determine highest degree
    degree_order = {"Master": 3, "Bachelor": 2, "Associate": 1, "Diploma": 0}
    highest_degree = max(
        education_entries,
        key=lambda e: degree_order.get(e["degree"], 0),
    )["degree"] if education_entries else None

    # Build resume_data dict (14 fields)
    resume_data: dict[str, Any] = {
        # Required fields (3)
        "full_name": profile.full_name_en,
        "work_experience": work_experiences,
        "total_positions": len(work_experiences),
        # Optional fields (11)
        "email": profile.contact_email,
        "phone": profile.contact_phone,
        "location": f"{profile.address.get('city', '')}, {profile.address.get('emirate', 'UAE')}".strip(", "),
        "summary": _generate_summary(profile.occupation, years_of_experience, rng),
        "years_of_experience": years_of_experience,
        "current_employer": profile.employer_name,
        "current_job_title": profile.occupation,
        "education": education_entries,
        "highest_degree": highest_degree,
        "skills": skills,
        "skill_count": skill_count,
        "certifications": certifications if certifications else None,
    }

    # Generate DOCX file
    output_dir = Path(__file__).parent.parent.parent / "output" / f"applicant_{seed}"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "resume.docx"

    _build_resume_docx(
        resume_data,
        work_experiences,
        education_entries,
        skills,
        certifications,
        output_path,
    )

    return resume_data, work_experiences, output_path
